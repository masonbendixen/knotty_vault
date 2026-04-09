"""Convert the Streaming Course Instruction Platform markdown to a Word docx."""

import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

INPUT_FILE = r"C:\Users\mason\Documents\Obsidian\Knotty Yoga\Claude\Streaming Course Instruction Platform.md"
OUTPUT_FILE = r"C:\Users\mason\Documents\Obsidian\Knotty Yoga\Claude\Streaming Course Instruction Platform.docx"


def strip_frontmatter(text):
    """Remove YAML frontmatter between --- delimiters."""
    if text.startswith("---"):
        end = text.find("---", 3)
        if end != -1:
            text = text[end + 3:].lstrip("\n")
    return text


def add_formatted_text(paragraph, text):
    """Parse inline markdown (bold, italic, code) and add runs to a paragraph."""
    # Pattern matches **bold**, *italic*, `code`, and plain text
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|([^*`]+))')
    for match in pattern.finditer(text):
        if match.group(2):  # **bold**
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif match.group(3):  # *italic*
            run = paragraph.add_run(match.group(3))
            run.italic = True
        elif match.group(4):  # `code`
            run = paragraph.add_run(match.group(4))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x80, 0x20, 0x20)
        elif match.group(5):  # plain text
            paragraph.add_run(match.group(5))


def parse_table(lines, start_idx):
    """Parse a markdown table starting at start_idx. Returns (rows, end_idx)."""
    rows = []
    i = start_idx
    while i < len(lines) and lines[i].strip().startswith("|"):
        row_text = lines[i].strip()
        # Skip separator rows (|---|---|)
        if re.match(r'^\|[\s\-:|]+\|$', row_text):
            i += 1
            continue
        cells = [c.strip() for c in row_text.split("|")[1:-1]]
        rows.append(cells)
        i += 1
    return rows, i


def add_table_to_doc(doc, rows):
    """Add a formatted table to the document."""
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = table.cell(i, j)
                cell.text = ""
                p = cell.paragraphs[0]
                p.style = doc.styles["Normal"]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                add_formatted_text(p, cell_text)
                # Bold the header row
                if i == 0:
                    for run in p.runs:
                        run.bold = True

    # Set font size for all cells
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph("")  # spacing after table


def convert_md_to_docx(input_path, output_path):
    with open(input_path, "r", encoding="utf-8") as f:
        content = f.read()

    content = strip_frontmatter(content)
    lines = content.split("\n")

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Adjust heading styles
    for level in range(1, 5):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code blocks
        if stripped.startswith("```"):
            if in_code_block:
                # End code block
                code_text = "\n".join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Horizontal rules
        if stripped == "---":
            # Add a subtle separator
            doc.add_paragraph("")
            i += 1
            continue

        # Headings
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped.lstrip("#").strip()
            level = min(level, 4)
            heading = doc.add_heading(level=level)
            add_formatted_text(heading, text)
            i += 1
            continue

        # Tables
        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            rows, end_idx = parse_table(lines, i)
            add_table_to_doc(doc, rows)
            i = end_idx
            continue

        # Bullet points (including nested)
        bullet_match = re.match(r'^(\s*)-\s+(.*)', line)
        if bullet_match:
            indent_level = len(bullet_match.group(1)) // 2
            text = bullet_match.group(2)
            p = doc.add_paragraph(style="List Bullet")
            p.clear()
            add_formatted_text(p, text)
            if indent_level > 0:
                p.paragraph_format.left_indent = Cm(1.27 * indent_level)
            i += 1
            continue

        # Numbered lists
        num_match = re.match(r'^(\s*)\d+\.\s+(.*)', line)
        if num_match:
            indent_level = len(num_match.group(1)) // 3
            text = num_match.group(2)
            p = doc.add_paragraph(style="List Number")
            p.clear()
            add_formatted_text(p, text)
            if indent_level > 0:
                p.paragraph_format.left_indent = Cm(1.27 * indent_level)
            i += 1
            continue

        # Empty lines
        if not stripped:
            i += 1
            continue

        # Regular paragraphs
        p = doc.add_paragraph()
        add_formatted_text(p, stripped)
        i += 1

    doc.save(output_path)
    print(f"Saved to: {output_path}")


if __name__ == "__main__":
    convert_md_to_docx(INPUT_FILE, OUTPUT_FILE)
