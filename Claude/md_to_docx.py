#!/usr/bin/env python3
"""
Convert Markdown to Word docx format.
Preserves mermaid code blocks as-is for later processing.
"""

import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path


def create_styles(doc):
    """Create custom styles for the document."""
    styles = doc.styles

    # Code block style
    if 'Code Block' not in [s.name for s in styles]:
        code_style = styles.add_style('Code Block', WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = 'Consolas'
        code_style.font.size = Pt(9)
        code_style.paragraph_format.left_indent = Inches(0.25)
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)


def parse_markdown(md_content):
    """Parse markdown content into structured elements."""
    lines = md_content.split('\n')
    elements = []
    i = 0

    # Skip YAML front matter
    if lines[0].strip() == '---':
        i = 1
        while i < len(lines) and lines[i].strip() != '---':
            i += 1
        i += 1  # Skip closing ---

    in_code_block = False
    code_block_content = []
    code_block_lang = ''
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]

        # Code blocks
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_block_lang = line.strip()[3:]
                code_block_content = []
            else:
                elements.append({
                    'type': 'code_block',
                    'lang': code_block_lang,
                    'content': '\n'.join(code_block_content)
                })
                in_code_block = False
            i += 1
            continue

        if in_code_block:
            code_block_content.append(line)
            i += 1
            continue

        # Tables
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []

            # Skip separator rows (|---|---|)
            if re.match(r'^\|[\s\-:|]+\|$', line.strip()):
                i += 1
                continue

            # Parse table row
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            # End of table
            elements.append({'type': 'table', 'rows': table_rows})
            in_table = False

        # Horizontal rule
        if line.strip() in ['---', '***', '___']:
            elements.append({'type': 'hr'})
            i += 1
            continue

        # Headers
        header_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if header_match:
            level = len(header_match.group(1))
            text = header_match.group(2)
            elements.append({'type': 'heading', 'level': level, 'text': text})
            i += 1
            continue

        # Bullet lists
        if re.match(r'^[\s]*[-*]\s+', line):
            indent = len(line) - len(line.lstrip())
            text = re.sub(r'^[\s]*[-*]\s+', '', line)
            elements.append({'type': 'bullet', 'text': text, 'indent': indent})
            i += 1
            continue

        # Numbered lists
        num_match = re.match(r'^[\s]*(\d+)\.\s+(.+)$', line)
        if num_match:
            indent = len(line) - len(line.lstrip())
            text = num_match.group(2)
            elements.append({'type': 'numbered', 'text': text, 'indent': indent})
            i += 1
            continue

        # Empty lines
        if not line.strip():
            elements.append({'type': 'empty'})
            i += 1
            continue

        # Regular paragraph
        elements.append({'type': 'paragraph', 'text': line})
        i += 1

    # Handle any remaining table
    if in_table:
        elements.append({'type': 'table', 'rows': table_rows})

    return elements


def apply_inline_formatting(paragraph, text):
    """Apply inline formatting (bold, italic, code) to text."""
    # Pattern to match **bold**, *italic*, `code`, and [links](url)
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\)|[^*`\[]+)'

    parts = re.findall(pattern, text)
    if not parts:
        paragraph.add_run(text)
        return

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        elif part.startswith('['):
            # Link - extract text and URL
            link_match = re.match(r'\[([^\]]+)\]\(([^)]+)\)', part)
            if link_match:
                link_text = link_match.group(1)
                link_url = link_match.group(2)
                run = paragraph.add_run(f"{link_text} ({link_url})")
                run.italic = True
        else:
            paragraph.add_run(part)


def convert_to_docx(elements, output_path):
    """Convert parsed elements to Word document."""
    doc = Document()
    create_styles(doc)

    for elem in elements:
        if elem['type'] == 'heading':
            level = elem['level']
            # Map markdown heading levels to Word heading styles
            if level == 1:
                p = doc.add_heading(elem['text'], level=0)  # Title
            else:
                p = doc.add_heading(elem['text'], level=min(level, 9))

        elif elem['type'] == 'paragraph':
            p = doc.add_paragraph()
            apply_inline_formatting(p, elem['text'])

        elif elem['type'] == 'bullet':
            p = doc.add_paragraph(style='List Bullet')
            apply_inline_formatting(p, elem['text'])
            if elem['indent'] > 0:
                p.paragraph_format.left_indent = Inches(elem['indent'] / 20)

        elif elem['type'] == 'numbered':
            p = doc.add_paragraph(style='List Number')
            apply_inline_formatting(p, elem['text'])
            if elem['indent'] > 0:
                p.paragraph_format.left_indent = Inches(elem['indent'] / 20)

        elif elem['type'] == 'code_block':
            # Add language label if present
            if elem['lang']:
                p = doc.add_paragraph()
                run = p.add_run(f"[{elem['lang']}]")
                run.font.size = Pt(8)
                run.italic = True

            # Add code content
            for line in elem['content'].split('\n'):
                p = doc.add_paragraph(line)
                p.style = doc.styles['Code Block'] if 'Code Block' in [s.name for s in doc.styles] else 'Normal'
                p.paragraph_format.left_indent = Inches(0.25)
                for run in p.runs:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)

        elif elem['type'] == 'table':
            rows = elem['rows']
            if rows:
                num_cols = max(len(row) for row in rows)
                table = doc.add_table(rows=len(rows), cols=num_cols)
                table.style = 'Table Grid'

                for i, row_data in enumerate(rows):
                    row = table.rows[i]
                    for j, cell_text in enumerate(row_data):
                        if j < len(row.cells):
                            cell = row.cells[j]
                            cell.text = ''
                            p = cell.paragraphs[0]
                            apply_inline_formatting(p, cell_text)
                            # Bold the header row
                            if i == 0:
                                for run in p.runs:
                                    run.bold = True

                doc.add_paragraph()  # Add space after table

        elif elem['type'] == 'hr':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run('─' * 50)
            run.font.color.rgb = None  # Use default color

        elif elem['type'] == 'empty':
            pass  # Skip empty lines (Word handles spacing)

    doc.save(output_path)
    print(f"Document saved to: {output_path}")


def main():
    input_path = Path(r"C:\Users\mason\Documents\Obsidian\Knotty Yoga\Claude\Payment Design Document.md")
    output_path = input_path.with_suffix('.docx')

    print(f"Reading: {input_path}")
    md_content = input_path.read_text(encoding='utf-8')

    print("Parsing markdown...")
    elements = parse_markdown(md_content)

    print(f"Converting to Word ({len(elements)} elements)...")
    convert_to_docx(elements, output_path)

    print("Done!")


if __name__ == '__main__':
    main()
